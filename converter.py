import mimetypes
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract
import os

class FileConverter:
    def __init__(self):
        mimetypes.init()

    def convert_to_docx(self, input_file):
        mime_type, _ = mimetypes.guess_type(input_file)
        output_file = os.path.splitext(input_file)[0] + ".docx"

        if mime_type == 'text/plain':
            self.txt_to_docx(input_file, output_file)
        elif mime_type == 'application/pdf':
            self.pdf_to_docx(input_file, output_file)
        elif mime_type in ['image/jpeg', 'image/png', 'image/gif']:
            self.image_to_docx(input_file, output_file)
        elif mime_type == 'application/msword':  # For .doc files
            self.doc_to_docx(input_file)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

        return output_file

    def txt_to_docx(self, txt_file, docx_file):
        doc = Document()
        with open(txt_file, 'r') as file:
            doc.add_paragraph(file.read())
        doc.save(docx_file)

    def pdf_to_docx(self, pdf_file, docx_file):
        doc = Document()
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
        doc.save(docx_file)

    def image_to_docx(self, image_file, docx_file):
        doc = Document()
        text = pytesseract.image_to_string(Image.open(image_file))
        doc.add_paragraph(text)
        doc.save(docx_file)

    # This doc_to_docx method is only for Windows with Word installed
    # def doc_to_docx(self, doc_file):
    #     import win32com.client as win32
    #     word = win32.gencache.EnsureDispatch('Word.Application')
    #     doc = word.Documents.Open(doc_file)
    #     doc.Activate()
    #     new_file_abs = doc_file.replace(".doc", ".docx")
    #     word.ActiveDocument.SaveAs(new_file_abs, FileFormat=16)
    #     doc.Close(False)

    
    def doc_to_docx(self, doc_file):
        import subprocess

        new_file_abs = doc_file.replace(".doc", ".docx")
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'docx', '--outdir', new_file_abs, doc_file], check=True)
       
converter = FileConverter()
output_docx = converter.convert_to_docx('МЕТРОЛОГІЧНА  ПЕРЕВІРКА  ЕЛЕКТРОННОГО ВОЛЬТМЕТРА (1).doc')  
print(f"Converted to DOCX: {output_docx}")
